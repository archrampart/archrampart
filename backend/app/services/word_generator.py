"""
Word (.docx) report generator for audit reports
Creates professional audit reports with cover page following international standards
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.orm import Session
from app.models.audit import Audit
from app.models.template import Severity, Status
from datetime import datetime
import tempfile
from collections import Counter

# Severity colors (RGB)
SEVERITY_COLORS = {
    Severity.CRITICAL: RGBColor(220, 53, 69),  # Red
    Severity.HIGH: RGBColor(253, 126, 20),     # Orange
    Severity.MEDIUM: RGBColor(255, 193, 7),    # Yellow
    Severity.LOW: RGBColor(13, 202, 240),      # Cyan
    Severity.INFO: RGBColor(108, 117, 125),    # Gray
}

# Status colors
STATUS_COLORS = {
    Status.OPEN: RGBColor(220, 53, 69),
    Status.IN_PROGRESS: RGBColor(255, 193, 7),
    Status.RESOLVED: RGBColor(22, 197, 94),
    Status.CLOSED: RGBColor(108, 117, 125),
}

def get_severity_text(severity: Severity) -> str:
    texts = {
        Severity.CRITICAL: "Kritik",
        Severity.HIGH: "Yüksek",
        Severity.MEDIUM: "Orta",
        Severity.LOW: "Düşük",
        Severity.INFO: "Bilgi"
    }
    return texts.get(severity, severity.value.upper())

def get_status_text(status: Status) -> str:
    texts = {
        Status.OPEN: "Açık",
        Status.IN_PROGRESS: "Devam Ediyor",
        Status.RESOLVED: "Çözüldü",
        Status.CLOSED: "Kapatıldı"
    }
    return texts.get(status, status.value.upper())

def add_cover_page(doc: Document, audit: Audit, organization, project):
    """Add professional cover page"""
    # Cover page section
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21.0)   # A4 width
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("DENETİM RAPORU")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue
    title_para.space_after = Pt(24)
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Audit name
    audit_name_para = doc.add_paragraph()
    audit_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audit_name_run = audit_name_para.add_run(audit.name)
    audit_name_run.font.size = Pt(20)
    audit_name_run.font.bold = True
    audit_name_para.space_after = Pt(12)
    
    # Standard
    if audit.standard:
        standard_para = doc.add_paragraph()
        standard_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        standard_run = standard_para.add_run(f"Standart: {audit.standard.value}")
        standard_run.font.size = Pt(14)
        standard_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Add vertical spacing
    for _ in range(8):
        doc.add_paragraph()
    
    # Information table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table content
    info_data = [
        ("Organizasyon", organization.name),
        ("Proje", project.name),
        ("Denetim Tarihi", audit.audit_date.strftime("%d.%m.%Y") if audit.audit_date else "Belirtilmemiş"),
        ("Rapor Tarihi", datetime.now().strftime("%d.%m.%Y")),
        ("Standart", audit.standard.value),
    ]
    
    for i, (label, value) in enumerate(info_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        
        # Style label cell
        label_cell = table.rows[i].cells[0]
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_para.runs[0]
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        
        # Style value cell
        value_cell = table.rows[i].cells[1]
        value_para = value_cell.paragraphs[0]
        value_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_run = value_para.runs[0]
        value_run.font.size = Pt(11)
    
    # Add page break
    doc.add_page_break()

def add_executive_summary(doc: Document, audit: Audit, findings):
    """Add executive summary section"""
    heading = doc.add_heading('ÖZET YÖNETİCİ RAPORU', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Summary paragraph
    summary_para = doc.add_paragraph()
    summary_para.add_run(
        f"Bu rapor, {audit.name} denetimi kapsamında yapılan incelemelerin sonuçlarını içermektedir. "
        f"Denetim {audit.standard.value} standartına uygun olarak gerçekleştirilmiş ve "
        f"toplamda {len(findings)} bulgu tespit edilmiştir."
    ).font.size = Pt(11)
    
    # Statistics
    severity_counts = Counter([f.severity for f in findings])
    
    doc.add_paragraph()
    stats_para = doc.add_paragraph("Bulgu Dağılımı:", style='List Bullet')
    stats_para.runs[0].font.bold = True
    stats_para.runs[0].font.size = Pt(11)
    
    for severity in [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]:
        count = severity_counts.get(severity, 0)
        if count > 0:
            stat_para = doc.add_paragraph(f"  • {get_severity_text(severity)}: {count} bulgu", style='List Bullet 2')
            stat_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()

def add_table_of_contents(doc: Document):
    """Add table of contents (placeholder)"""
    heading = doc.add_heading('İÇİNDEKİLER', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    toc_items = [
        "1. Özet Yönetici Raporu",
        "2. Denetim Kapsamı ve Metodolojisi",
        "3. Bulgular",
        "4. Sonuç ve Öneriler",
        "5. Ekler"
    ]
    
    for item in toc_items:
        toc_para = doc.add_paragraph(item)
        toc_para.runs[0].font.size = Pt(11)
        toc_para.space_after = Pt(6)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_page_break()

def add_methodology_section(doc: Document, audit: Audit):
    """Add methodology and scope section"""
    heading = doc.add_heading('DENETİM KAPSAMI VE METODOLOJİSİ', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Scope
    scope_para = doc.add_paragraph()
    scope_para.add_run("Denetim Kapsamı:").font.bold = True
    scope_para.add_run().font.size = Pt(11)
    
    scope_content = doc.add_paragraph(f"Bu denetim {audit.standard.value} standartı kapsamında gerçekleştirilmiştir.")
    scope_content.runs[0].font.size = Pt(11)
    
    if audit.description:
        doc.add_paragraph()
        desc_para = doc.add_paragraph()
        desc_para.add_run("Açıklama:").font.bold = True
        desc_para.add_run().font.size = Pt(11)
        
        desc_content = doc.add_paragraph(audit.description)
        desc_content.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()

def add_findings_section(doc: Document, findings):
    """Add detailed findings section"""
    heading = doc.add_heading('BULGULAR', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if not findings:
        no_findings = doc.add_paragraph("Denetim kapsamında bulgu tespit edilmemiştir.")
        no_findings.runs[0].font.size = Pt(11)
        no_findings.runs[0].italic = True
        return
    
    for idx, finding in enumerate(findings, 1):
        # Finding heading
        finding_heading = doc.add_heading(f"BULGU #{idx}: {finding.title}", level=2)
        finding_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Severity and Status badges
        badge_para = doc.add_paragraph()
        badge_para.paragraph_format.space_after = Pt(6)
        
        # Severity badge
        severity_run = badge_para.add_run(f"Önem Derecesi: {get_severity_text(finding.severity)}")
        severity_run.font.bold = True
        severity_run.font.size = Pt(10)
        severity_run.font.color.rgb = SEVERITY_COLORS.get(finding.severity, RGBColor(0, 0, 0))
        severity_run.add_text("  |  ")
        
        # Status badge
        status_run = badge_para.add_run(f"Durum: {get_status_text(finding.status)}")
        status_run.font.bold = True
        status_run.font.size = Pt(10)
        status_run.font.color.rgb = STATUS_COLORS.get(finding.status, RGBColor(0, 0, 0))
        
        # Control reference
        if finding.control_reference:
            ref_para = doc.add_paragraph()
            ref_para.add_run("Kontrol Referansı: ").font.bold = True
            ref_para.add_run(finding.control_reference).font.size = Pt(11)
        
        # Description
        if finding.description:
            desc_para = doc.add_paragraph()
            desc_para.add_run("Açıklama:").font.bold = True
            desc_para.add_run().font.size = Pt(11)
            
            desc_content = doc.add_paragraph(finding.description)
            desc_content.runs[0].font.size = Pt(11)
            desc_content.paragraph_format.space_after = Pt(6)
        
        # Recommendation
        if finding.recommendation:
            rec_para = doc.add_paragraph()
            rec_para.add_run("Öneri:").font.bold = True
            rec_para.add_run().font.size = Pt(11)
            
            rec_content = doc.add_paragraph(finding.recommendation)
            rec_content.runs[0].font.size = Pt(11)
            rec_content.paragraph_format.space_after = Pt(6)
        
        # Evidence count
        if finding.evidences:
            evid_para = doc.add_paragraph()
            evid_para.add_run(f"Kanıt Sayısı: {len(finding.evidences)}").font.bold = True
            evid_para.add_run().font.size = Pt(11)
            
            for evidence in finding.evidences:
                evid_item = doc.add_paragraph(f"  • {evidence.file_name}", style='List Bullet')
                evid_item.runs[0].font.size = Pt(10)
                if evidence.description:
                    evid_desc = doc.add_paragraph(f"    Açıklama: {evidence.description}", style='List Bullet 2')
                    evid_desc.runs[0].font.size = Pt(9)
                    evid_desc.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        
        # Add spacing between findings
        doc.add_paragraph()
        doc.add_paragraph()

def add_conclusion_section(doc: Document, audit: Audit, findings):
    """Add conclusion and recommendations section"""
    heading = doc.add_heading('SONUÇ VE ÖNERİLER', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run(
        f"Bu denetim raporu, {audit.name} denetimi kapsamında tespit edilen bulguları ve "
        f"önerileri içermektedir. Organizasyonun ilgili standartlara uygunluğunu artırmak için "
        f"belirtilen önerilerin değerlendirilmesi ve uygulanması önerilmektedir."
    ).font.size = Pt(11)
    
    doc.add_paragraph()

def add_appendix_section(doc: Document, audit: Audit):
    """Add appendix section"""
    heading = doc.add_heading('EKLER', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    appendix_para = doc.add_paragraph()
    appendix_para.add_run("Ek A: Denetim Metadataları").font.bold = True
    appendix_para.add_run().font.size = Pt(11)
    
    metadata_items = [
        ("Denetim ID", str(audit.id)),
        ("Oluşturulma Tarihi", audit.created_at.strftime("%d.%m.%Y %H:%M") if audit.created_at else "N/A"),
        ("Son Güncelleme", audit.updated_at.strftime("%d.%m.%Y %H:%M") if audit.updated_at else "N/A"),
    ]
    
    for label, value in metadata_items:
        meta_para = doc.add_paragraph(f"  {label}: {value}")
        meta_para.runs[0].font.size = Pt(10)

def generate_audit_word_report(audit: Audit, db: Session) -> str:
    """
    Generate professional Word (.docx) report for an audit
    Follows international audit report standards with cover page
    """
    
    # Get audit details with relationships
    findings = audit.findings
    project = audit.project
    organization = project.organization
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = f"Denetim Raporu - {audit.name}"
    doc.core_properties.author = "ArchRampart Audit Tool"
    doc.core_properties.comments = f"{audit.standard.value} Denetim Raporu"
    
    # Add cover page
    add_cover_page(doc, audit, organization, project)
    
    # Add table of contents
    add_table_of_contents(doc)
    
    # Add executive summary
    add_executive_summary(doc, audit, findings)
    
    # Add methodology
    add_methodology_section(doc, audit)
    
    # Add findings
    add_findings_section(doc, findings)
    
    # Add conclusion
    add_conclusion_section(doc, audit, findings)
    
    # Add appendix
    add_appendix_section(doc, audit)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_path = temp_file.name
    temp_file.close()
    
    doc.save(temp_path)
    
    return temp_path

